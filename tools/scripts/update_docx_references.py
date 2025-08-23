import datetime
import shutil
from pathlib import Path

from docx import Document


def append_references(docx_path: Path, references: list[str]) -> None:
    # Backup first
    backup = docx_path.with_suffix(docx_path.suffix + ".bak")
    try:
        shutil.copy2(docx_path, backup)
    except Exception:
        pass

    doc = Document(str(docx_path))

    # Simple idempotency: if a recent section exists, still add a new dated section
    today = datetime.date.today().isoformat()
    doc.add_paragraph("")
    doc.add_heading(f"Repository References ({today})", level=2)
    for ref in references:
        # Use plain paragraph with a leading bullet for maximum compatibility
        doc.add_paragraph(f"• {ref}")

    doc.save(str(docx_path))


def main():
    repo = Path(__file__).resolve().parents[2]
    doc_dir = repo / "assets" / ".md.txt"
    docx_files = sorted(doc_dir.glob("*.docx"))
    if not docx_files:
        print("No .docx files found under assets/.md.txt")
        return

    # Curated references to newly added or updated files
    refs = [
        "orchestration/server.py — Agent Bridge API (/search, /plan, /health)",
        "orchestration/openapi.yaml — OpenAPI spec for ChatGPT Actions",
        "orchestration/.env.example — DATACORE_URL and PORT",
        ".vscode/tasks.json — tasks to run Datacore and Orchestration",
        "docs/guides/TEC_Multi_Agent_Creative_System.md — multi‑agent playbook",
        "docs/REPO_OVERHAUL_PLAN.md — repo cleanup and Drive sync guidance",
        "SESSION_HANDOFF_CHAT_ARCHIVE.md — session snapshot and next steps",
    ]

    for f in docx_files:
        append_references(f, refs)
        print(f"Updated: {f}")


if __name__ == "__main__":
    main()
